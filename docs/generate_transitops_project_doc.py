from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(__file__).with_name('TransitOps_Complete_Project_Document.docx')

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

styles = doc.styles
styles['Normal'].font.name = 'Aptos'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
styles['Normal'].font.size = Pt(10)
for name, size, color in [('Title', 28, '0B1F3A'), ('Heading 1', 18, '0B1F3A'), ('Heading 2', 13, '007A68'), ('Heading 3', 11, '374151')]:
    styles[name].font.name = 'Aptos Display'
    styles[name]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos Display')
    styles[name].font.size = Pt(size)
    styles[name].font.color.rgb = RGBColor.from_string(color)

# Custom compact style
if 'Small Note' not in styles:
    s = styles.add_style('Small Note', WD_STYLE_TYPE.PARAGRAPH)
    s.font.name = 'Aptos'
    s.font.size = Pt(8)
    s.font.color.rgb = RGBColor(90, 98, 108)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, True, 'FFFFFF')
        shade(t.rows[0].cells[i], '0B1F3A')
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if len(t.rows) % 2 == 0:
                shade(cells[i], 'F2F7F7')
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()
    return t


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)
    return p


def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)
    return p


def callout(title, text, fill='E8F5F2'):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shade(cell, fill)
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(title + '\n')
    r.bold = True
    r.font.color.rgb = RGBColor(0, 92, 80)
    p.add_run(text)
    doc.add_paragraph()

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(18)
r = p.add_run('TRANSITOPS')
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = RGBColor(0, 122, 104)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Complete Project Document')
r.bold = True
r.font.size = Pt(25)
r.font.color.rgb = RGBColor(11, 31, 58)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Smart Fleet Operations Platform with Zero Trust Continuous Authorization')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(55, 65, 81)

doc.add_paragraph('\n')
callout('PROJECT PURPOSE', 'A complete design and implementation record for fleet operations, request-level security, policy enforcement, continuous audit, simulation, and the ZTCA Security Admin Panel.', 'DDEFF4')

table(['Document control', 'Value'], [
    ('Project', 'TransitOps Smart Fleet Operations Platform'),
    ('Security model', 'Zero Trust Continuous Authorization (ZTCA)'),
    ('Primary stack', 'React, Vite, TypeScript, Express, JSON repositories'),
    ('Primary audience', 'Project reviewers, administrators, developers, operators'),
    ('Document status', 'Complete implementation and operating guide'),
])
doc.add_page_break()

# Contents

doc.add_heading('Contents', level=1)
for item in [
    '1. Project Overview', '2. Problem Statement and Objectives', '3. Requirements and Scope',
    '4. Top-to-Bottom Project Making Process', '5. System Design and Architecture',
    '6. Seven Zero Trust Cyber Policies', '7. End-to-End Operational Workflows',
    '8. User Interface and Security Admin Panel', '9. Data Model and Persistence',
    '10. Testing and Demonstration Cases', '11. Deployment and Operations',
    '12. Risks, Controls, and Future Roadmap', '13. Conclusion', 'Appendix A. API and File Map', 'Appendix B. Test Checklist'
]:
    doc.add_paragraph(item, style='List Bullet')
doc.add_page_break()

# 1
doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph('TransitOps is a full-stack fleet operations platform for managing vehicles, drivers, trips, maintenance, expenses, reports, and operational activity. The platform adds a Zero Trust Continuous Authorization layer so every protected API request is evaluated using user identity, role, device fingerprint, location, time window, endpoint sensitivity, and active policy rules.')
doc.add_paragraph('The project combines operational efficiency with security observability. The ZTCA Security Admin Panel provides a live request stream, outcome metrics, risk-factor inspection, policy management, device registry management, and decision analytics.')
table(['Layer', 'Responsibility', 'Project implementation'], [
    ('Presentation', 'User workflows and security visualization', 'React components in src/components and App.tsx'),
    ('API', 'Business routes and request handling', 'Express routes in backend/routes.ts'),
    ('Authorization', 'Risk calculation and PDP decisions', 'backend/ztca/engine.ts and middleware.ts'),
    ('Persistence', 'Operational and security records', 'JSON repositories and backend/data'),
    ('Build and run', 'Development and production delivery', 'Vite, TypeScript, tsx, esbuild'),
])

# 2
doc.add_heading('2. Problem Statement and Objectives', level=1)
doc.add_heading('2.1 Problem statement', level=2)
doc.add_paragraph('Fleet operations contain high-impact actions. Dispatching a route, changing a driver status, modifying a vehicle, recording financial activity, or changing a security policy can affect safety, cost, compliance, and service continuity. Login-only security is insufficient because the context of a request can change after authentication.')
doc.add_heading('2.2 Objectives', level=2)
for x in [
    'Evaluate access continuously at request level instead of trusting a session indefinitely.',
    'Use risk scoring from 0 to 100 and explain each contributing factor.',
    'Return clear decisions: ALLOW, STEP_UP, READ_ONLY, or BLOCK.',
    'Record both successful and denied activity in an auditable stream.',
    'Allow administrators to test simulations without modifying operational records.',
    'Reflect 100/100 simulated risk as an immediate BLOCK and increment blocked metrics.',
    'Give fleet managers usable workflows without hiding security decisions.'
]: bullet(x)

# 3
doc.add_heading('3. Requirements and Scope', level=1)
table(['Area', 'In scope'], [
    ('Fleet operations', 'Vehicles, drivers, trips, maintenance, expenses, fuel, reports'),
    ('Identity', 'User role, user profile, login and registration flows'),
    ('ZTCA context', 'User, role, device, browser, OS, location, time, endpoint, action'),
    ('PDP outcomes', 'Allow, Step-Up MFA, Read-Only, Block'),
    ('Administration', 'Policies, devices, locations, audit logs, metrics, inspection'),
    ('Simulation', 'Context presets, risk preview, forced 100/100 block, audit reflection'),
    ('Persistence', 'Repository-backed JSON data suitable for local/demo operation'),
])
doc.add_heading('3.1 Out of scope for the demonstration build', level=2)
for x in ['Production identity provider integration', 'Distributed database replication', 'Hardware-backed device attestation', 'Formal legal/compliance certification', 'Production-grade event streaming at scale']: bullet(x)

# 4
doc.add_heading('4. Top-to-Bottom Project Making Process', level=1)
doc.add_paragraph('The project is built in a deliberate sequence so business workflows, security decisions, user experience, and evidence remain connected.')
steps = [
    ('Step 1: Discover', 'Identify fleet roles, actions, resources, risks, and operational failure modes.'),
    ('Step 2: Define roles and resources', 'Model Admin, Fleet Manager, Driver, Mechanic, and other roles against vehicles, drivers, trips, finance, and security administration.'),
    ('Step 3: Design the data model', 'Define operational entities plus ZTCA context, risk, policy, decision, device, location, and audit structures.'),
    ('Step 4: Build repositories', 'Create focused repository classes for users, vehicles, drivers, trips, maintenance, expenses, devices, locations, policies, and audit records.'),
    ('Step 5: Build API routes', 'Expose CRUD and operational routes through Express, then centralize protected request evaluation in middleware.'),
    ('Step 6: Build the ZTCA engine', 'Calculate risk factors, aggregate a bounded score, classify risk level, evaluate policies, and produce a decision.'),
    ('Step 7: Enforce decisions', 'Allow safe requests, issue MFA challenges, downgrade writes to read-only, or return a blocked response with audit id.'),
    ('Step 8: Build the UI', 'Create operational dashboards and the ZTCA Admin Panel with live stream, metrics, charts, inspector, policies, and devices.'),
    ('Step 9: Add simulation', 'Provide trusted, unknown-device, foreign-location, off-hours, and high-risk presets for repeatable security tests.'),
    ('Step 10: Connect reflection', 'Persist simulation and administrative events, update metrics, label simulations, and broadcast immediate KPI updates.'),
    ('Step 11: Validate', 'Run happy-path, MFA, read-only, blocked, suspension, policy-change, and 100/100 simulation cases.'),
    ('Step 12: Package and operate', 'Build with Vite and esbuild, protect environment secrets, retain logs, and plan migration from JSON to a database.'),
]
table(['Process stage', 'Work performed'], steps, [2.0, 4.8])

# 5
doc.add_heading('5. System Design and Architecture', level=1)
doc.add_heading('5.1 Architecture', level=2)
doc.add_paragraph('The browser communicates with the Express application. Protected API traffic passes through the ZTCA authorization middleware before route handlers execute. The middleware constructs context from request headers, consults trust repositories, evaluates the engine, persists the audit record, and enforces the decision.')
callout('REQUEST FLOW', 'Browser action -> Express API -> ZTCA context collection -> risk evaluation -> policy decision -> audit persistence -> enforcement -> Admin Panel reflection.', 'DDEFF4')
doc.add_heading('5.2 Core components', level=2)
table(['Component', 'Purpose'], [
    ('React application', 'Operational screens and simulation controls'),
    ('ZTCAContextWidget', 'Context presets, dry-run risk forecast, factor preview'),
    ('AdminPanel', 'Live security stream, KPI cards, analytics, policy and device administration'),
    ('ztcaAuthorizationMiddleware', 'Interception, context construction, audit logging, enforcement'),
    ('ZTCAEngine', 'Risk scoring and policy decision point'),
    ('ZTCAAuditRepository', 'Persistent ordered audit events, maximum 500 records in demo mode'),
    ('JSONStore', 'Simple file-backed repository abstraction'),
])
doc.add_heading('5.3 Decision states', level=2)
table(['Decision', 'Meaning', 'Typical response'], [
    ('ALLOW', 'Context satisfies current controls', 'Request proceeds'),
    ('STEP_UP', 'Risk is elevated but can be mitigated by MFA', '401 with verification challenge'),
    ('READ_ONLY', 'Read access is acceptable but write action is unsafe', 'Write denied; read access remains conceptually available'),
    ('BLOCK', 'Action is forbidden by role, policy, critical risk, or simulation rule', '403 with reason and auditLogId'),
])

# 6
doc.add_heading('6. Seven Zero Trust Cyber Policies', level=1)
doc.add_paragraph('The following seven policies provide a complete project-aligned Zero Trust control set. Policies 1 to 5 are represented by the configured policy repository. Policies 6 and 7 formalize controls already implemented through the middleware, audit handlers, and simulation reflection flow.')
policies = [
    ('Policy 1: Verify Every Request', 'Never rely only on a successful login. Construct a fresh context for every protected API request.', 'Middleware reads user, role, device, location, time, endpoint, method, and action headers.', 'All protected /api routes', 'No request is trusted automatically because it belongs to an authenticated session.'),
    ('Policy 2: Enforce Least Privilege', 'A user may perform only actions consistent with their role and required privilege.', 'Admin routes are restricted to Admin; driver elevated operations create privilege mismatch risk; endpoint sensitivity is evaluated.', 'Admin, driver, vehicle, expense, and trip modifications', 'Unprivileged attempts are blocked or escalated depending on policy and risk.'),
    ('Policy 3: Trust Devices and Locations Explicitly', 'Device fingerprints and operating locations must be known and trusted before sensitive work is accepted.', 'Known-device and known-location repositories are checked for each request.', 'All sensitive writes and policy changes', 'Unknown device adds risk; foreign or unregistered location adds risk; policy may require trusted context.'),
    ('Policy 4: Apply Adaptive Risk-Based Access', 'Access level must change according to context risk, not remain static.', 'Engine adds weighted factors and classifies LOW, MEDIUM, HIGH, or CRITICAL risk.', 'Every protected request', 'Low-risk requests may be allowed; moderate writes can become read-only; high-risk requests require MFA or block.'),
    ('Policy 5: Protect Sensitive Operations', 'Dispatch, fleet changes, financial changes, and security administration receive stronger controls.', 'Endpoint and HTTP method sensitivity rules add risk and configured policy thresholds apply action-specific controls.', 'Trips, vehicles, drivers, expenses, and admin policies', 'Sensitive operations are allowed only when role, context, and risk conditions pass.'),
    ('Policy 6: Continuously Audit and Reflect Decisions', 'Every decision, including allowed, denied, MFA, read-only, simulated, and administrative changes, must be traceable.', 'ZTCAAuditRepository stores context, risk factors, decision, actor, endpoint, timestamp, and audit id. Admin Panel polls logs and metrics.', 'All ZTCA decisions and security administration', 'The live stream and inspector show who acted, what happened, why it happened, and the resulting outcome.'),
    ('Policy 7: Fail Closed for Critical Simulation and Account Actions', 'A full-risk simulation or account restriction must be treated as blocked and visible in security metrics.', 'A simulation with risk 100 is forced to BLOCK and tagged SIMULATION_FORCED_BLOCK. Driver/user status restrictions and device/policy changes are logged.', '100/100 simulations, suspension/block actions, trust changes', 'Blocked count increments, the row is labeled Simulated where applicable, and the Admin Panel displays the event.'),
]
for title, intent, implementation, scope, outcome in policies:
    doc.add_heading(title, level=2)
    table(['Control item', 'Description'], [('Security intent', intent), ('Implementation in TransitOps', implementation), ('Protected scope', scope), ('Expected outcome', outcome)], [1.7, 5.1])

# 7
doc.add_heading('7. End-to-End Operational Workflows', level=1)
doc.add_heading('7.1 Normal allowed operation', level=2)
for x in ['User selects an operational action.', 'Frontend sends the request with ZTCA context headers.', 'Middleware verifies trust records and calculates risk.', 'PDP returns ALLOW.', 'Audit entry is persisted and route handler changes operational data.', 'Admin Panel reflects the allowed request in the stream and metrics.']: numbered(x)
doc.add_heading('7.2 Step-Up MFA workflow', level=2)
for x in ['A Fleet Manager attempts a sensitive dispatch from elevated-risk context.', 'ZTCA returns STEP_UP with a verification challenge and auditLogId.', 'User submits the verification PIN through the step-up endpoint.', 'The client retries the action with the temporary step-up token.', 'ZTCA recognizes the verified token and permits the action if all remaining policy conditions pass.', 'Both the challenge and successful retry remain visible in the audit stream.']: numbered(x)
doc.add_heading('7.3 Read-only workflow', level=2)
for x in ['A write request reaches a moderate-risk threshold.', 'PDP returns READ_ONLY.', 'Middleware rejects the modification and returns a read-only error with auditLogId.', 'The event contributes to the read-only count and remains inspectable with its risk factors.']: numbered(x)
doc.add_heading('7.4 Blocked account or unprivileged action', level=2)
for x in ['A Driver attempts an elevated operation or a restricted account is suspended.', 'Role, account, policy, or critical-risk controls identify the action as forbidden.', 'The system returns BLOCK or records a blocked administrative status change.', 'The audit record includes actor, target action, reason, context, and outcome.', 'Admin Panel increments Blocked and displays the event in the Live Request Stream.']: numbered(x)
doc.add_heading('7.5 Simulation reflection workflow', level=2)
for x in ['Administrator changes a simulation preset or context field.', 'Simulation endpoint evaluates the context without changing fleet data.', 'A simulation audit event is persisted with simulation=true.', 'At risk 100, the decision is forced to BLOCK with SIMULATION_FORCED_BLOCK.', 'The endpoint returns an updated metrics snapshot.', 'The widget broadcasts the snapshot and Admin Panel updates Blocked immediately.']: numbered(x)

# 8
doc.add_heading('8. User Interface and Security Admin Panel', level=1)
doc.add_heading('8.1 Operational interface', level=2)
for x in ['Dashboard for operational summary', 'Vehicle registry and profile changes', 'Driver roster, compliance, and suspension status', 'Trip dispatch and status updates', 'Maintenance, expenses, fuel, and reports', 'ZTCA context widget for simulation and risk preview']: bullet(x)
doc.add_heading('8.2 Security Admin Panel', level=2)
table(['Panel area', 'Displayed information'], [
    ('KPI row', 'Total Requests, Allowed, Step-Up MFA, Read-Only, Blocked, Average Context Risk'),
    ('Live Request Stream', 'Timestamp, user and role, action and route, checkups, risk score, decision, inspect action'),
    ('Inspector drawer', 'Decision reason, timestamp, device, location, risk factors, policy trigger, simulation marker'),
    ('Decision Analytics', 'Outcome distribution, risk-level distribution, device, location, and off-hours checkups'),
    ('Policy Configurator', 'Enable/disable rules, risk threshold, action if violated, custom policy creation'),
    ('Device Registry', 'Fingerprint, browser, OS, last active time, trust status, trust/revoke operations'),
])
doc.add_heading('8.3 Reflection rules', level=2)
for x in ['Simulation rows show a Simulated marker.', 'Blocked simulation decisions show an explicit SIMULATION FORCED BLOCK reason.', 'Blocked KPI is calculated from persisted audit decisions, including simulations.', 'Administrative driver/device/user/policy actions are logged with actor information when available.', 'The Admin Panel refreshes through DB polling and receives an immediate simulation metrics broadcast.']: bullet(x)

# 9
doc.add_heading('9. Data Model and Persistence', level=1)
doc.add_heading('9.1 ZTCA audit record', level=2)
table(['Field group', 'Important fields'], [
    ('Actor', 'id, name, email, role'),
    ('Context', 'device, browser, OS, known flags, city, country, coordinates, odd-hours flag'),
    ('Request', 'endpoint, method, actionName, requiredPrivilege'),
    ('Risk', 'totalScore, level, factors with rule id, name, score, description'),
    ('Decision', 'outcome, riskScore, riskLevel, reason, policyTriggered, stepUpChallenge, timestamp'),
    ('Reflection', 'simulation marker and generated auditLogId'),
])
doc.add_heading('9.2 Demo persistence behavior', level=2)
for x in ['Operational data is stored in backend/data JSON files.', 'ZTCA policies, devices, locations, and audit logs are persisted separately.', 'Audit entries are sorted newest first and retained up to 500 entries.', 'The JSON repository is appropriate for local demonstration and testing; migrate to a transactional database for production.']: bullet(x)

# 10
doc.add_heading('10. Testing and Demonstration Cases', level=1)
table(['Case', 'Test context', 'Expected result in Admin Panel'], [
    ('Allow', 'Known user, trusted device, verified location, normal hours, low-risk read', 'ALLOW row and Allowed count increment'),
    ('Step-Up MFA', 'Fleet Manager dispatch from unknown device or off-hours', 'STEP_UP row, challenge visible, Step-Up count increment'),
    ('Read-Only', 'Moderate-risk write request above read-only threshold', 'READ_ONLY row, reason and factors visible'),
    ('Unprivileged block', 'Driver attempts admin or restricted operation', 'BLOCK row, actor and privilege mismatch visible'),
    ('Suspension action', 'Manager changes driver status to Suspended', 'Administrative audit row identifies actor and blocked status action'),
    ('100/100 simulation', 'Unknown device + foreign location + odd-hours + sensitive endpoint + privilege mismatch', 'Forced BLOCK, Simulated marker, Blocked count increment immediately'),
    ('Policy change', 'Admin enables, disables, creates, or updates a rule', 'Audit row identifies policy action and acting administrator'),
])
doc.add_heading('10.1 100/100 simulation acceptance criteria', level=2)
for x in ['API response contains risk.totalScore = 100.', 'API response contains decision.outcome = BLOCK.', 'Decision reason begins with SIMULATION FORCED BLOCK.', 'Response contains auditLogId.', 'Persisted audit record has context.simulation = true and decision.simulation = true.', 'GET /api/admin/metrics includes the new record in blockedCount.', 'Admin Panel shows the blocked row and the Blocked KPI increases.']: bullet(x)

# 11
doc.add_heading('11. Deployment and Operations', level=1)
doc.add_heading('11.1 Development', level=2)
for x in ['Install dependencies with npm install.', 'Use npm run dev to start the Express plus Vite development server on port 3000.', 'Keep backend/data files out of the Vite watch path so data updates do not cause continuous browser refreshes.', 'Use the Admin Panel and simulation widget for live validation.']: numbered(x)
doc.add_heading('11.2 Production build', level=2)
for x in ['Run npm run build.', 'Vite creates the client build and esbuild bundles server.ts into dist/server.cjs.', 'Run npm start to start the bundled server.', 'Protect server-side secrets such as GEMINI_API_KEY in environment configuration; do not expose them in browser code.']: numbered(x)
doc.add_heading('11.3 Operational controls', level=2)
for x in ['Review blocked and step-up trends regularly.', 'Investigate repeated unknown-device or location anomalies.', 'Review policy threshold changes through the audit stream.', 'Retain and export audit logs according to organizational retention requirements.', 'Replace JSON persistence with a database before multi-user deployment.']: bullet(x)

# 12
doc.add_heading('12. Risks, Controls, and Future Roadmap', level=1)
table(['Risk', 'Current control', 'Future improvement'], [
    ('JSON file concurrency', 'Repository abstraction and small local workload', 'Transactional database with locking'),
    ('Header-supplied identity', 'ZTCA context headers and role checks in demo', 'Signed tokens and identity provider integration'),
    ('Polling overhead', '1.5-second polling plus simulation broadcast', 'SSE or WebSocket event stream'),
    ('Simulation/live confusion', 'Simulation marker and explicit forced-block reason', 'Separate simulation view and report filter'),
    ('Secret exposure', 'Gemini key intended for server-side use', 'Secret manager, rotation, usage monitoring'),
    ('False positives', 'Step-Up MFA and trusted device/location registry', 'Adaptive baselines and analyst review'),
])
doc.add_heading('12.1 Roadmap', level=2)
for x in ['Add simulation-only and live-only filters.', 'Open the Admin Panel inspector directly from a simulation result.', 'Add SSE/WebSocket updates for audit events and metrics.', 'Add database migrations and audit retention controls.', 'Integrate secure server-side Gemini summaries for audit trends.', 'Add automated integration tests for all seven policies.']: bullet(x)

# 13
doc.add_heading('13. Conclusion', level=1)
doc.add_paragraph('TransitOps demonstrates how a fleet platform can combine business operations with continuous, context-aware authorization. The seven-policy model makes the security posture understandable: verify every request, enforce least privilege, trust devices and locations explicitly, adapt access to risk, protect sensitive operations, audit every decision, and fail closed for critical simulations and account actions.')
doc.add_paragraph('The most important operational result is reflection. A simulation is not complete merely because it calculates a score. When the score reaches 100/100, the system must block it, persist it, count it, and show it in the Security Admin Panel. The same principle applies to MFA, read-only decisions, suspensions, policy changes, and other security events.')

# Appendix A
doc.add_heading('Appendix A. API and File Map', level=1)
table(['Area', 'Location'], [
    ('Application shell', 'src/App.tsx'),
    ('Admin security UI', 'src/components/AdminPanel.tsx'),
    ('Simulation widget', 'src/components/ZTCAContextWidget.tsx'),
    ('ZTCA middleware', 'backend/ztca/middleware.ts'),
    ('Risk and decision engine', 'backend/ztca/engine.ts'),
    ('ZTCA data types', 'backend/ztca/types.ts'),
    ('ZTCA repositories', 'backend/ztca/store.ts'),
    ('API routes', 'backend/routes.ts'),
    ('Audit data', 'backend/data/ztcaAuditLogs.json'),
    ('Policy data', 'backend/data/ztcaPolicies.json'),
    ('Development configuration', 'vite.config.ts'),
])
doc.add_heading('Key endpoints', level=2)
table(['Endpoint', 'Purpose'], [
    ('POST /api/ztca/simulation/context-check', 'Evaluate and log a simulated context'),
    ('POST /api/ztca/verify-stepup', 'Verify demo step-up PIN'),
    ('GET /api/admin/audit-logs', 'Read filtered audit stream'),
    ('GET /api/admin/metrics', 'Read current outcome and risk metrics'),
    ('GET/POST/PUT/DELETE /api/admin/policies', 'Manage policy rules'),
    ('GET/PUT /api/admin/devices', 'Read and update device trust records'),
])

# Appendix B
doc.add_heading('Appendix B. Test Checklist', level=1)
for x in [
    'Start the application and confirm the Admin Panel loads.',
    'Confirm a normal trusted simulation returns ALLOW.',
    'Confirm unknown device and foreign location contribute risk factors.',
    'Confirm high-risk dispatch returns STEP_UP where configured.',
    'Confirm a moderate-risk write returns READ_ONLY where configured.',
    'Confirm a Driver cannot access Admin policy endpoints.',
    'Confirm 100/100 simulation returns BLOCK and SIMULATION FORCED BLOCK.',
    'Confirm the 100/100 audit row appears in Live Request Stream.',
    'Confirm Blocked KPI increases after the 100/100 simulation.',
    'Confirm inspector shows score, factor details, actor, reason, and simulation marker.',
    'Confirm driver suspension records who performed the action.',
    'Confirm policy and device changes appear in the audit stream.',
    'Confirm Vite does not refresh the browser continuously when backend JSON changes.',
    'Confirm npm run build completes before release.'
]:
    bullet('[ ] ' + x)

# Footer
for sec in doc.sections:
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('TransitOps Complete Project Document | ZTCA Security Architecture').font.size = Pt(8)

doc.core_properties.title = 'TransitOps Complete Project Document'
doc.core_properties.subject = 'Seven Zero Trust cyber policies, project design, workflow, and implementation'
doc.core_properties.author = 'TransitOps Project Team'
doc.save(OUT)
print(OUT)
